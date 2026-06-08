"""
PDF crash report extractor.

Extracts structured carrier/vehicle fields from text-based PDFs using
label-pattern regex. One record is produced per detected vehicle unit;
if no unit boundaries are found, one record is produced per page.
"""

import io
import re
import uuid
import pdfplumber
import pandas as pd

# ---------------------------------------------------------------------------
# Field-level regex patterns (case-insensitive, applied to uppercase page text)
# ---------------------------------------------------------------------------
_PATTERNS: dict[str, list[str]] = {
    'crash_report_number': [
        r'(?:CRASH|ACCIDENT)\s+REPORT\s+(?:NUMBER|NO\.?|#)\s*[:\-]?\s*([A-Z0-9\-]+)',
        r'REPORT\s*(?:NUMBER|NO\.?|#)\s*[:\-]?\s*([A-Z0-9\-]+)',
        r'REPORT\s*#\s*([A-Z0-9\-]+)',
    ],
    'case_number': [
        r'CASE\s*(?:NUMBER|NO\.?|#)\s*[:\-]?\s*([A-Z0-9\-]+)',
    ],
    'crash_date': [
        r'(?:DATE\s+OF\s+(?:CRASH|ACCIDENT)|CRASH\s+DATE|ACCIDENT\s+DATE)\s*[:\-]?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
        r'\bDATE\s*[:\-]\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
    ],
    'vehicle_number': [
        r'(?:UNIT|VEHICLE)\s*(?:NUMBER|NO\.?|#)\s*[:\-]?\s*([0-9]+)',
        r'\bUNIT\s*[:\-]?\s*#?\s*([0-9]+)',
        r'\bVEH(?:ICLE)?\s*[:\-]?\s*#?\s*([0-9]+)',
    ],
    'driver_name': [
        r'(?:DRIVER|OPERATOR)\s*(?:NAME)?\s*[:\-]\s*([A-Z][A-Z\s,\.]+?)(?=\s{2,}|\n|DOB|DATE|LIC)',
        r'\bDRIVER\s*[:\-]\s*([^\n]+)',
    ],
    'driver_address': [
        r'(?:DRIVER|OPERATOR)\s+ADDRESS\s*[:\-]\s*([^\n]+)',
        r'DRV\s+ADDR(?:ESS)?\s*[:\-]\s*([^\n]+)',
    ],
    'carrier_name': [
        r'(?:MOTOR\s+)?CARRIER\s+(?:NAME\s*)?[:\-]\s*([^\n]+)',
        r'COMPANY\s+NAME\s*[:\-]\s*([^\n]+)',
        r'TRUCKING\s+(?:COMPANY\s*)?[:\-]\s*([^\n]+)',
    ],
    'carrier_address': [
        r'CARRIER\s+ADDR(?:ESS)?\s*[:\-]\s*([^\n]+)',
        r'COMPANY\s+ADDR(?:ESS)?\s*[:\-]\s*([^\n]+)',
    ],
    'owner_name': [
        r'(?:REGISTERED\s+)?OWNER\s+(?:NAME\s*)?[:\-]\s*([^\n]+?)(?=\s{2,}|\n|ADDR)',
        r'\bOWNER\s*[:\-]\s*([^\n]+)',
    ],
    'owner_address': [
        r'(?:REGISTERED\s+)?OWNER\s+ADDR(?:ESS)?\s*[:\-]\s*([^\n]+)',
        r'OWNER\s+ADDR\s*[:\-]\s*([^\n]+)',
    ],
    'usdot_number': [
        r'(?:US\s*)?DOT\s*(?:NUMBER|NO\.?|#)?\s*[:\-]?\s*#?\s*(\d{5,9})',
        r'USDOT\s*[:\-]?\s*#?\s*(\d{5,9})',
    ],
    'vin': [
        r'\bVIN\s*[:\-]\s*([A-HJ-NPR-Z0-9]{17})\b',
        r'\b([A-HJ-NPR-Z0-9]{17})\b',
    ],
    'vehicle_plate': [
        r'(?:LICENSE|PLATE|TAG)\s*(?:PLATE\s*)?(?:NUMBER|NO\.?|#)?\s*[:\-]\s*([A-Z0-9]{2,10})',
        r'\bPLATE\s*[:\-]\s*([A-Z0-9]{2,10})',
        r'\bTAG\s*[:\-]\s*([A-Z0-9]{2,10})',
    ],
    'vehicle_plate_state': [
        r'(?:LICENSE|PLATE)\s+ST(?:ATE)?\s*[:\-]\s*([A-Z]{2})\b',
        r'PLATE\s+ST\s*[:\-]\s*([A-Z]{2})\b',
    ],
    'vehicle_make': [
        r'\bMAKE\s*[:\-]\s*([A-Z]{2,20})\b',
        r'VEH(?:ICLE)?\s+MAKE\s*[:\-]\s*([A-Z]{2,20})\b',
    ],
    'vehicle_year': [
        r'\bYEAR\s*[:\-]\s*((?:19|20)\d{2})\b',
        r'VEH(?:ICLE)?\s+YEAR\s*[:\-]\s*((?:19|20)\d{2})\b',
    ],
    'insurance_company': [
        r'INS(?:URANCE)?\s*(?:COMPANY|CO\.?)?\s*[:\-]\s*([^\n]+)',
        r'INSURER\s*[:\-]\s*([^\n]+)',
    ],
}

_UNIT_SPLIT = re.compile(
    r'(?:^|\n)\s*(?:UNIT|VEHICLE)\s*(?:NUMBER\s*)?[:\-]?\s*#?\s*[0-9]+',
    re.IGNORECASE,
)


def _match_field(text_upper: str, field: str) -> str:
    for pat in _PATTERNS.get(field, []):
        try:
            m = re.search(pat, text_upper, re.IGNORECASE | re.MULTILINE)
            if m:
                return m.group(1).strip()
        except Exception:
            pass
    return ''


def _extract_zip(address: str) -> str:
    m = re.search(r'\b(\d{5})(?:-\d{4})?\b', address)
    return m.group(1) if m else ''


def _parse_unit(text: str, page_num: int, source: str, unit_idx: int) -> dict:
    t = text.upper()
    carrier_addr = _match_field(t, 'carrier_address')
    owner_addr = _match_field(t, 'owner_address')

    return {
        'pdf_record_id': f"{source[:6]}-p{page_num}-u{unit_idx}-{uuid.uuid4().hex[:4]}",
        'source_file': source,
        'pdf_page_number': page_num,
        'crash_report_number': _match_field(t, 'crash_report_number'),
        'case_number': _match_field(t, 'case_number'),
        'crash_date': _match_field(t, 'crash_date'),
        'vehicle_number': _match_field(t, 'vehicle_number') or str(unit_idx),
        'driver_name': _match_field(t, 'driver_name'),
        'driver_address': _match_field(t, 'driver_address'),
        'carrier_name': _match_field(t, 'carrier_name'),
        'carrier_address_raw': carrier_addr,
        'carrier_zip': _extract_zip(carrier_addr),
        'owner_name': _match_field(t, 'owner_name'),
        'owner_address_raw': owner_addr,
        'owner_zip': _extract_zip(owner_addr),
        'usdot_number': _match_field(t, 'usdot_number'),
        'vin': _match_field(t, 'vin'),
        'vehicle_plate_number': _match_field(t, 'vehicle_plate'),
        'vehicle_plate_state': _match_field(t, 'vehicle_plate_state'),
        'vehicle_make': _match_field(t, 'vehicle_make'),
        'vehicle_year': _match_field(t, 'vehicle_year'),
        'insurance_company': _match_field(t, 'insurance_company'),
        'raw_text': text[:600],
    }


def _has_data(rec: dict) -> bool:
    key_fields = ['crash_report_number', 'carrier_name', 'carrier_address_raw',
                  'owner_name', 'usdot_number', 'vin', 'vehicle_plate_number']
    return any(rec.get(f, '').strip() for f in key_fields)


def extract_pdf_records(file, source_filename: str = 'unknown.pdf') -> pd.DataFrame:
    """
    Open a PDF and return a DataFrame of extracted vehicle/carrier records.

    Falls back to one record per page when no unit-boundary markers are detected.
    """
    records: list[dict] = []

    try:
        raw = file.read() if hasattr(file, 'read') else open(file, 'rb').read()
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ''
                if not text.strip():
                    continue

                # Try to split by UNIT / VEHICLE markers
                splits = _UNIT_SPLIT.split(text)
                if len(splits) > 1:
                    # First split is preamble (crash header), rest are unit blocks
                    header = splits[0]
                    for idx, block in enumerate(splits[1:], 1):
                        combined = header + '\n' + block
                        rec = _parse_unit(combined, page_num, source_filename, idx)
                        records.append(rec)
                else:
                    rec = _parse_unit(text, page_num, source_filename, 1)
                    if _has_data(rec) or page_num == 1:
                        records.append(rec)

    except Exception as exc:
        records.append({
            'pdf_record_id': 'err-' + uuid.uuid4().hex[:6],
            'source_file': source_filename,
            'pdf_page_number': 0,
            'crash_report_number': '',
            'case_number': '',
            'crash_date': '',
            'vehicle_number': '',
            'driver_name': '',
            'driver_address': '',
            'carrier_name': '',
            'carrier_address_raw': '',
            'carrier_zip': '',
            'owner_name': '',
            'owner_address_raw': '',
            'owner_zip': '',
            'usdot_number': '',
            'vin': '',
            'vehicle_plate_number': '',
            'vehicle_plate_state': '',
            'vehicle_make': '',
            'vehicle_year': '',
            'insurance_company': '',
            'raw_text': f'Extraction error: {exc}',
        })

    return pd.DataFrame(records) if records else pd.DataFrame()


def _text_to_records(text: str, source_filename: str) -> list[dict]:
    """Run the same unit-split + field-extraction logic on a plain text string."""
    records = []
    splits = _UNIT_SPLIT.split(text)
    if len(splits) > 1:
        header = splits[0]
        for idx, block in enumerate(splits[1:], 1):
            records.append(_parse_unit(header + '\n' + block, 1, source_filename, idx))
    else:
        rec = _parse_unit(text, 1, source_filename, 1)
        records.append(rec)
    return records


def extract_docx_records(file_bytes: bytes, source_filename: str) -> pd.DataFrame:
    """
    Extract crash-report fields from a Word .docx file.
    Uses python-docx to pull all paragraph text, then applies the same
    pattern logic as the PDF extractor.
    """
    try:
        import docx as _docx
        import io as _io
        doc = _docx.Document(_io.BytesIO(file_bytes))
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        # Also pull table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += '\n' + cell.text
        records = _text_to_records(full_text, source_filename)
        return pd.DataFrame(records) if records else pd.DataFrame()
    except Exception as exc:
        return pd.DataFrame([{
            'pdf_record_id': 'err-' + uuid.uuid4().hex[:6],
            'source_file': source_filename,
            'pdf_page_number': 0,
            **{k: '' for k in [
                'crash_report_number', 'case_number', 'crash_date', 'vehicle_number',
                'driver_name', 'driver_address', 'carrier_name', 'carrier_address_raw',
                'carrier_zip', 'owner_name', 'owner_address_raw', 'owner_zip',
                'usdot_number', 'vin', 'vehicle_plate_number', 'vehicle_plate_state',
                'vehicle_make', 'vehicle_year', 'insurance_company',
            ]},
            'raw_text': f'DOCX extraction error: {exc}',
        }])


def extract_plaintext_records(file_bytes: bytes, source_filename: str) -> pd.DataFrame:
    """
    Extract crash-report fields from a plain-text file (Google Doc export,
    .txt, .rtf stripped content, etc.).
    """
    try:
        text = file_bytes.decode('utf-8', errors='replace')
        records = _text_to_records(text, source_filename)
        return pd.DataFrame(records) if records else pd.DataFrame()
    except Exception as exc:
        return pd.DataFrame([{
            'pdf_record_id': 'err-' + uuid.uuid4().hex[:6],
            'source_file': source_filename,
            'pdf_page_number': 0,
            **{k: '' for k in [
                'crash_report_number', 'case_number', 'crash_date', 'vehicle_number',
                'driver_name', 'driver_address', 'carrier_name', 'carrier_address_raw',
                'carrier_zip', 'owner_name', 'owner_address_raw', 'owner_zip',
                'usdot_number', 'vin', 'vehicle_plate_number', 'vehicle_plate_state',
                'vehicle_make', 'vehicle_year', 'insurance_company',
            ]},
            'raw_text': f'Text extraction error: {exc}',
        }])
